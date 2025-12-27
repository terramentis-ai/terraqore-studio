"""
Flynt Context Mining Tool
Extracts ideas and context from various file formats and sources.
Supports: PDF, TXT, MD, DOCX, URLs, and more.
"""

import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
import re

logger = logging.getLogger(__name__)


@dataclass
class ExtractedContext:
    """Extracted context from a source."""
    source: str  # File path or URL
    source_type: str  # pdf, txt, url, docx, etc.
    content: str
    metadata: Dict[str, Any]
    word_count: int
    success: bool = True
    error: Optional[str] = None


class ContextMiner:
    """Mines context and ideas from various sources."""
    
    def __init__(self):
        """Initialize context miner."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check availability of optional dependencies."""
        self.deps = {
            'pypdf': self._check_import('pypdf'),
            'docx': self._check_import('docx'),
            'youtube_transcript_api': self._check_import('youtube_transcript_api'),
            'requests': self._check_import('requests'),
            'beautifulsoup4': self._check_import('bs4')
        }
        
        # Log available capabilities
        available = [k for k, v in self.deps.items() if v]
        if available:
            logger.info(f"Context mining available for: {', '.join(available)}")
        else:
            logger.warning("No optional dependencies installed. Only basic text/markdown supported.")
    
    def _check_import(self, module_name: str) -> bool:
        """Check if a module can be imported."""
        try:
            if module_name == 'docx':
                __import__('docx')
            elif module_name == 'bs4':
                __import__('bs4')
            else:
                __import__(module_name)
            return True
        except ImportError:
            return False
    
    def mine_file(self, file_path: str) -> ExtractedContext:
        """Mine context from a file.
        
        Args:
            file_path: Path to file.
            
        Returns:
            ExtractedContext object.
        """
        path = Path(file_path)
        
        if not path.exists():
            return ExtractedContext(
                source=file_path,
                source_type="unknown",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=f"File not found: {file_path}"
            )
        
        # Determine file type and extract
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf(path)
        elif suffix in ['.txt', '.md', '.markdown']:
            return self._extract_text(path)
        elif suffix in ['.docx', '.doc']:
            return self._extract_docx(path)
        else:
            # Try as text
            return self._extract_text(path)
    
    def mine_url(self, url: str) -> ExtractedContext:
        """Mine context from a URL.
        
        Args:
            url: URL to extract from.
            
        Returns:
            ExtractedContext object.
        """
        # Check if YouTube URL
        if 'youtube.com' in url or 'youtu.be' in url:
            return self._extract_youtube(url)
        else:
            return self._extract_webpage(url)
    
    def _extract_text(self, path: Path) -> ExtractedContext:
        """Extract text from plain text or markdown file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            word_count = len(content.split())
            
            return ExtractedContext(
                source=str(path),
                source_type=path.suffix[1:],
                content=content,
                metadata={'file_size': path.stat().st_size},
                word_count=word_count,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ExtractedContext(
                source=str(path),
                source_type="text",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=str(e)
            )
    
    def _extract_pdf(self, path: Path) -> ExtractedContext:
        """Extract text from PDF file."""
        if not self.deps['pypdf']:
            return ExtractedContext(
                source=str(path),
                source_type="pdf",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error="pypdf not installed. Run: pip install pypdf"
            )
        
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(path)
            content = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
            
            full_content = "\n\n".join(content)
            word_count = len(full_content.split())
            
            return ExtractedContext(
                source=str(path),
                source_type="pdf",
                content=full_content,
                metadata={
                    'pages': len(reader.pages),
                    'file_size': path.stat().st_size
                },
                word_count=word_count,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ExtractedContext(
                source=str(path),
                source_type="pdf",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=str(e)
            )
    
    def _extract_docx(self, path: Path) -> ExtractedContext:
        """Extract text from Word document."""
        if not self.deps['docx']:
            return ExtractedContext(
                source=str(path),
                source_type="docx",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error="python-docx not installed. Run: pip install python-docx"
            )
        
        try:
            from docx import Document
            
            doc = Document(path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            full_content = "\n\n".join(content)
            word_count = len(full_content.split())
            
            return ExtractedContext(
                source=str(path),
                source_type="docx",
                content=full_content,
                metadata={
                    'paragraphs': len(doc.paragraphs),
                    'file_size': path.stat().st_size
                },
                word_count=word_count,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ExtractedContext(
                source=str(path),
                source_type="docx",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=str(e)
            )
    
    def _extract_youtube(self, url: str) -> ExtractedContext:
        """Extract transcript from YouTube video."""
        if not self.deps['youtube_transcript_api']:
            return ExtractedContext(
                source=url,
                source_type="youtube",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error="youtube-transcript-api not installed. Run: pip install youtube-transcript-api"
            )
        
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            
            # Extract video ID
            video_id = None
            if 'v=' in url:
                video_id = url.split('v=')[1].split('&')[0]
            elif 'youtu.be/' in url:
                video_id = url.split('youtu.be/')[1].split('?')[0]
            
            if not video_id:
                raise ValueError("Could not extract video ID from URL")
            
            # Get transcript
            transcript = YouTubeTranscriptApi.get_transcript(video_id)
            
            # Combine transcript text
            content = " ".join([entry['text'] for entry in transcript])
            word_count = len(content.split())
            
            return ExtractedContext(
                source=url,
                source_type="youtube",
                content=content,
                metadata={
                    'video_id': video_id,
                    'transcript_length': len(transcript)
                },
                word_count=word_count,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error extracting YouTube transcript: {e}")
            return ExtractedContext(
                source=url,
                source_type="youtube",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=f"Could not extract transcript: {str(e)}"
            )
    
    def _extract_webpage(self, url: str) -> ExtractedContext:
        """Extract main content from webpage."""
        if not self.deps['requests'] or not self.deps['beautifulsoup4']:
            return ExtractedContext(
                source=url,
                source_type="webpage",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error="requests and beautifulsoup4 not installed. Run: pip install requests beautifulsoup4"
            )
        
        try:
            import requests
            from bs4 import BeautifulSoup
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style', 'nav', 'footer', 'header']):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            word_count = len(content.split())
            
            # Try to get title
            title = soup.find('title')
            title_text = title.string if title else "Unknown"
            
            return ExtractedContext(
                source=url,
                source_type="webpage",
                content=content,
                metadata={'title': title_text},
                word_count=word_count,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Error extracting webpage: {e}")
            return ExtractedContext(
                source=url,
                source_type="webpage",
                content="",
                metadata={},
                word_count=0,
                success=False,
                error=str(e)
            )
    
    def mine_multiple(self, sources: List[str]) -> List[ExtractedContext]:
        """Mine context from multiple sources.
        
        Args:
            sources: List of file paths or URLs.
            
        Returns:
            List of ExtractedContext objects.
        """
        results = []
        
        for source in sources:
            if source.startswith('http://') or source.startswith('https://'):
                result = self.mine_url(source)
            else:
                result = self.mine_file(source)
            
            results.append(result)
        
        return results
    
    def extract_key_ideas(self, content: str, max_length: int = 2000) -> str:
        """Extract key ideas from long content.
        
        Args:
            content: Full content text.
            max_length: Maximum length to return.
            
        Returns:
            Truncated content focused on key ideas.
        """
        # Simple extraction: prioritize paragraphs with key idea indicators
        paragraphs = content.split('\n\n')
        
        key_indicators = [
            'key', 'important', 'main', 'crucial', 'essential',
            'summary', 'conclusion', 'insight', 'idea', 'concept',
            'therefore', 'thus', 'because', 'should', 'must'
        ]
        
        scored_paragraphs = []
        for para in paragraphs:
            if len(para.strip()) < 20:
                continue
            
            score = sum(1 for indicator in key_indicators if indicator in para.lower())
            scored_paragraphs.append((score, para))
        
        # Sort by score and take top paragraphs
        scored_paragraphs.sort(reverse=True, key=lambda x: x[0])
        
        result = []
        current_length = 0
        
        for score, para in scored_paragraphs:
            if current_length + len(para) > max_length:
                break
            result.append(para)
            current_length += len(para)
        
        if not result:
            # If no paragraphs matched, just take the beginning
            return content[:max_length]
        
        return '\n\n'.join(result)


# Singleton instance
_context_miner: Optional[ContextMiner] = None


def get_context_miner() -> ContextMiner:
    """Get or create the global context miner.
    
    Returns:
        ContextMiner singleton instance.
    """
    global _context_miner
    if _context_miner is None:
        _context_miner = ContextMiner()
    return _context_miner